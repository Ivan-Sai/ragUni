"""Build a reference.docx with the styles required by the faculty.

Pandoc uses ``--reference-doc`` as a template — the styles defined
in this docx (font, size, margins, headings, paragraph spacing) are
applied to every paragraph in the converted output. This script
generates that template programmatically so we don't depend on a
manually-edited Word file.

Requirements per the faculty's methodology:
  * Font: Times New Roman 14 pt
  * Line spacing: 1.5
  * Paragraph indent: 1.25 cm
  * Page size: A4
  * Margins: left 25 mm, right 15 mm, top 20 mm, bottom 20 mm
  * Heading 1: 14 pt, BOLD, CENTERED, UPPERCASE (no period)
  * Heading 2: 14 pt, bold, paragraph indent (NOT centered, no period)
  * Page numbers: top-right, no period (added in convert step)

Usage::

    python build_reference_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


def _set_default_font(doc: Document) -> None:
    """Set Times New Roman 14 pt as the document-wide default."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 0, 0)
    # Ensure Cyrillic glyphs use TNR (rFonts cs / eastAsia)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Paragraph format: line spacing 1.5, first-line indent 1.25 cm
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _configure_section(doc: Document) -> None:
    """A4 page size + faculty-required margins."""
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)


def _style_heading(style, *, centered: bool, bold: bool, upper: bool) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    if centered:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(1.25)
    if upper:
        # ``All caps`` runs at the rendered level; we toggle the
        # corresponding XML attribute so even if author's source isn't
        # uppercase, Word still renders headings in capitals.
        rPr_caps = rPr.find(qn("w:caps"))
        if rPr_caps is None:
            rPr.append(rPr.makeelement(qn("w:caps"), {qn("w:val"): "true"}))


def _add_page_numbers(doc: Document) -> None:
    """Top-right page numbers, no period after the digit."""
    section = doc.sections[0]
    header = section.header
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Build a PAGE field
    run = para.add_run()
    fld_begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    instr = run._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = " PAGE "
    fld_end = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def main() -> int:
    out = Path(__file__).resolve().parent / "reference.docx"
    doc = Document()
    _set_default_font(doc)
    _configure_section(doc)

    # H1: section titles (РОЗДІЛ 1, ВСТУП, ВИСНОВКИ, etc.)
    _style_heading(
        doc.styles["Heading 1"],
        centered=True, bold=True, upper=True,
    )
    # H2: subsections (1.1, 1.2, ...)
    _style_heading(
        doc.styles["Heading 2"],
        centered=False, bold=True, upper=False,
    )
    # H3: sub-subsections (1.4.1, 1.4.2, ...)
    _style_heading(
        doc.styles["Heading 3"],
        centered=False, bold=True, upper=False,
    )

    _add_page_numbers(doc)

    doc.save(out)
    print(f"Wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
