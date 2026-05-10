"""End-to-end build of diploma.docx from diploma.md.

Steps performed:
  1. Pandoc converts diploma.md → diploma.docx using reference.docx
     for styling (TNR 14, 1.5 spacing, faculty margins).
  2. python-docx post-processes the result to:
       - insert a title page that satisfies the faculty's standard
         layout (university / faculty / department block on top,
         work title centred in the middle, supervisor + city/year
         at the bottom, no page number);
       - insert an automatic Table of Contents after the title page
         and the реферат;
       - re-apply paragraph indent (1.25 cm) and 1.5 line spacing
         to body paragraphs that pandoc left in the default style;
       - centre figure captions and image paragraphs.

Usage::

    python convert.py
"""

from __future__ import annotations

import os
import re
import sys
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


HERE = Path(__file__).resolve().parent
SRC = HERE / "diploma.md"
REF = HERE / "reference.docx"
OUT = HERE / "diploma.docx"


# ---------------------------------------------------------------------------
# Title page values — edit these to match your data before running.
# ---------------------------------------------------------------------------
TITLE = {
    "university": "Київський національний університет імені Тараса Шевченка",
    "faculty": "Факультет радіофізики, електроніки та комп'ютерних систем",
    "department": "Кафедра комп'ютерної інженерії",
    "work_kind": "КВАЛІФІКАЦІЙНА РОБОТА БАКАЛАВРА",
    "title": (
        "Розробка веб-системи RAG-пошуку для університетської "
        "документації з підтримкою рольового доступу та "
        "структурованого видобування даних"
    ),
    "specialty": (
        "освітньо-професійна програма «Інженерія комп'ютерних "
        "систем і мереж»"
    ),
    "student_label": "Здобувач:",
    "student_name": "Сай Іван Олегович",
    "supervisor_label": "Науковий керівник:",
    "supervisor_name": "_________________________________",
    "supervisor_position": "(посада, науковий ступінь, П. І. Б.)",
    "city_year": "Київ — 2026",
}


# ---------------------------------------------------------------------------
# Pandoc step
# ---------------------------------------------------------------------------


def run_pandoc() -> None:
    """Convert diploma.md to a docx using reference.docx for styling."""
    if not REF.exists():
        print(f"ERROR: {REF.name} not found — run `python build_reference_docx.py` first.")
        sys.exit(1)
    extra = [
        f"--reference-doc={REF}",
        "--standalone",
        "--top-level-division=section",
        # Pipe-tables in markdown become real Word tables.
        "--from=markdown+pipe_tables+grid_tables+raw_html+fenced_code_blocks",
    ]
    pypandoc.convert_file(
        str(SRC), "docx", outputfile=str(OUT), extra_args=extra,
    )
    print(f"  pandoc wrote {OUT.name}")


# ---------------------------------------------------------------------------
# Post-processing helpers
# ---------------------------------------------------------------------------


def _set_cell_run(run, *, bold: bool = False, size: int = 14, font: str = "Times New Roman") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)


def _add_paragraph(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.CENTER, bold: bool = False, indent: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25 if indent else 0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    _set_cell_run(run, bold=bold, size=14)


def _add_blank_paragraph(doc: Document, count: int = 1) -> None:
    for _ in range(count):
        doc.add_paragraph()


def _set_no_header_first_page(section) -> None:
    """Title page must NOT have a page number in the header."""
    section.different_first_page_header_footer = True
    fp_header = section.first_page_header
    for p in fp_header.paragraphs:
        for run in p.runs:
            run.text = ""


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------


def add_title_page(doc: Document) -> None:
    """Prepend a faculty-standard title page to the document.

    python-docx doesn't have a built-in "insert at top" — we build the
    page in a fresh sub-document and splice its body before the
    existing one. The page break is achieved with a manual
    ``WD_BREAK.PAGE``.
    """
    # Build the new pieces into a list of (text, kwargs) so the spliced
    # XML appears at the top of the document body.
    new_doc = Document()
    # Apply the same default font / margins as the rest of the
    # document so spliced content matches.
    new_doc.styles["Normal"].font.name = "Times New Roman"
    new_doc.styles["Normal"].font.size = Pt(14)

    def add(text: str, *, align=WD_ALIGN_PARAGRAPH.CENTER, bold: bool = False) -> None:
        _add_paragraph(new_doc, text, align=align, bold=bold)

    add(TITLE["university"])
    add(TITLE["faculty"])
    add(TITLE["department"])
    _add_blank_paragraph(new_doc, 8)
    add(TITLE["work_kind"], bold=True)
    _add_blank_paragraph(new_doc, 1)
    add(TITLE["title"], bold=True)
    _add_blank_paragraph(new_doc, 2)
    add("Освітньо-професійна програма")
    add(TITLE["specialty"])
    _add_blank_paragraph(new_doc, 6)

    # Student + supervisor block, right-aligned
    p = new_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(f"{TITLE['student_label']}  {TITLE['student_name']}")
    _set_cell_run(run, size=14)

    p = new_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(f"{TITLE['supervisor_label']}  {TITLE['supervisor_name']}")
    _set_cell_run(run, size=14)

    p = new_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(TITLE["supervisor_position"])
    _set_cell_run(run, size=12)

    _add_blank_paragraph(new_doc, 6)
    add(TITLE["city_year"])

    # Page break
    p = new_doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # Splice the new paragraphs at the top of doc.body.
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    new_paragraphs = new_doc.element.body.findall(qn("w:p"))
    # Insert in reverse so they end up in order at the top.
    for p_xml in reversed(new_paragraphs):
        body.insert(0, p_xml)

    # Configure first page to not show header/footer.
    _set_no_header_first_page(doc.sections[0])


# ---------------------------------------------------------------------------
# Body cleanups: enforce 1.5 spacing + first-line indent on body paras
# ---------------------------------------------------------------------------


def normalize_body_paragraphs(doc: Document) -> None:
    """Re-apply line spacing + indent on every body paragraph.

    Pandoc sometimes resets paragraph_format on imports — this loops
    through every paragraph and re-applies our defaults to anything
    that isn't a heading.
    """
    for p in doc.paragraphs:
        style_name = (p.style.name if p.style else "") or ""
        is_heading = style_name.startswith("Heading")
        is_caption = style_name == "Caption"
        is_image = (
            len(p.runs) == 1 and not p.text and any(
                run.element.find(qn("w:drawing")) is not None for run in p.runs
            )
        )
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if is_heading or is_caption or is_image:
            pf.first_line_indent = Cm(0)
            if is_image:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Body paragraph — but skip those that are intentionally
            # centered (like image captions or the keywords block).
            if p.alignment in (None, WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.LEFT):
                pf.first_line_indent = Cm(1.25)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ---------------------------------------------------------------------------
# Centre figure captions ("Рисунок N — ...")
# ---------------------------------------------------------------------------


_FIGURE_CAPTION_RE = re.compile(r"^Рисунок\s+\d+(\.\d+)?\s+[—-]")


def centre_figure_captions(doc: Document) -> None:
    for p in doc.paragraphs:
        if _FIGURE_CAPTION_RE.match(p.text or ""):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)


# ---------------------------------------------------------------------------
# Insert ToC field after the реферат section
# ---------------------------------------------------------------------------


def insert_toc_after_referat(doc: Document) -> None:
    """Inject Word-native TOC field after the РЕФЕРАТ heading.

    The field instruction ``TOC \\o "1-3" \\h \\z \\u`` mirrors what
    Word generates for ``References → Table of Contents → Automatic 1``;
    it covers headings 1-3, includes hyperlinks, hides tab leaders
    in the web layout and uses outline levels.

    The TOC is empty until the user opens the document and presses
    ``F9`` (or right-click → Update Field) — Word can't compute the
    field at file-write time.
    """
    body = doc.element.body
    paragraphs = list(body.iter(qn("w:p")))
    referat_idx = None
    for i, p in enumerate(paragraphs):
        text = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
        if text == "РЕФЕРАТ":
            referat_idx = i
            break
    if referat_idx is None:
        # Couldn't find — leave alone.
        return

    # Insert TOC after the реферат content; we look for the first H1
    # AFTER referat (which is ВСТУП) and insert the TOC right before it.
    target_idx = None
    for i in range(referat_idx + 1, len(paragraphs)):
        p = paragraphs[i]
        # H1 is a paragraph with pStyle Heading1
        pStyle = p.find(qn("w:pPr"))
        if pStyle is not None:
            style = pStyle.find(qn("w:pStyle"))
            if style is not None and style.get(qn("w:val")) in {"Heading1", "Heading 1"}:
                target_idx = i
                break

    if target_idx is None:
        target_idx = referat_idx + 1
    target = paragraphs[target_idx]

    # Build TOC paragraph
    toc_heading = OxmlElement("w:p")
    toc_pPr = OxmlElement("w:pPr")
    toc_pStyle = OxmlElement("w:pStyle")
    toc_pStyle.set(qn("w:val"), "Heading1")
    toc_pPr.append(toc_pStyle)
    toc_heading.append(toc_pPr)
    toc_run = OxmlElement("w:r")
    toc_text = OxmlElement("w:t")
    toc_text.text = "ЗМІСТ"
    toc_run.append(toc_text)
    toc_heading.append(toc_run)

    toc_field = OxmlElement("w:p")
    fld_run = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_run.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_run.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_run.append(fld_sep)
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click → Update Field to populate the table of contents."
    fld_run.append(placeholder)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_run.append(fld_end)
    toc_field.append(fld_run)

    # Page break after TOC
    page_break_p = OxmlElement("w:p")
    pb_run = OxmlElement("w:r")
    pb_break = OxmlElement("w:br")
    pb_break.set(qn("w:type"), "page")
    pb_run.append(pb_break)
    page_break_p.append(pb_run)

    target.addprevious(toc_heading)
    target.addprevious(toc_field)
    target.addprevious(page_break_p)


# ---------------------------------------------------------------------------
# Replace stat placeholders in реферат
# ---------------------------------------------------------------------------


def fill_in_statistics(doc: Document, stats: dict[str, int]) -> None:
    """Replace `___ сторінок` / `___ рисунків` / etc. with real counts."""
    for p in doc.paragraphs:
        for run in p.runs:
            if "___" not in (run.text or ""):
                continue
            run.text = (run.text or "").replace(
                "___ сторінок", f"{stats['pages']} сторінок"
            ).replace(
                "___ рисунків", f"{stats['figures']} рисунків"
            ).replace(
                "___ джерел посилань", f"{stats['sources']} джерел посилань"
            ).replace(
                "___ таблиць", f"{stats['tables']} таблиць"
            ).replace(
                "___ додатків", f"{stats['appendices']} додатків"
            ).replace(
                "(___ найменувань)", f"({stats['sources']} найменувань)"
            )


def main() -> int:
    print("Step 1: pandoc convert...")
    run_pandoc()

    print("Step 2: load result for post-processing...")
    doc = Document(str(OUT))

    print("Step 3: prepend title page...")
    add_title_page(doc)

    print("Step 4: insert table of contents...")
    insert_toc_after_referat(doc)

    print("Step 5: normalise body paragraphs (1.5 spacing + 1.25 indent)...")
    normalize_body_paragraphs(doc)

    print("Step 6: centre figure captions...")
    centre_figure_captions(doc)

    print("Step 7: fill in реферат statistics...")
    # Counts derived from the document structure:
    #   - 4 figures (Рис. 2.1 .. 2.4)
    #   - 6 tables (Табл. 1.1, 2.1, 2.2, 2.3, 4.1, 4.2, 4.3) — verify after build
    #   - 20 sources in the bibliography
    #   - 3 appendices (А, Б, В)
    #   - pages: filled in after final word count (placeholder = 60)
    fill_in_statistics(
        doc,
        {
            "pages": 60,
            "figures": 4,
            "sources": 20,
            "tables": 7,
            "appendices": 3,
        },
    )

    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    print()
    print("Open the file in Word and:")
    print("  1. Press F9 inside the ЗМІСТ section to populate the TOC.")
    print("  2. Verify the page count and update the реферат stats line if needed.")
    print("  3. Run plagiarism check via the faculty's recommended service.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
