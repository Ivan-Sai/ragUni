from docx import Document

doc = Document('RAG_Report_Draft.docx')

# Знайти і замінити неточний текст про тестові дані
for paragraph in doc.paragraphs:
    if 'Завантажено графік сесії з 41 групою, методичні матеріали' in paragraph.text:
        # Видалити старий текст
        for run in paragraph.runs:
            run.text = ''

        # Додати правильний текст
        paragraph.add_run('Завантажено 2 документи: розклад занять та графік екзаменаційної сесії. Графік сесії містить розклад іспитів для 41 групи університету. Загалом створено близько 80 векторних chunks.')

print('✓ Виправлено секцію з тестовими даними')
doc.save('RAG_Report_Draft.docx')
print('✓ Документ оновлено з реальними цифрами')
