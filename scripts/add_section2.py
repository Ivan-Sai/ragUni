from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('RAG_Report_Draft.docx')

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_paragraph(doc, text, bold_start=None):
    p = doc.add_paragraph()
    if bold_start:
        p.add_run(bold_start).bold = True
        p.add_run(' ' + text)
    else:
        p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

# РОЗДІЛ 2
add_heading(doc, 'РОЗДІЛ 2', 1)
add_heading(doc, 'ПРОЕКТУВАННЯ СИСТЕМИ', 1)

add_heading(doc, '2.1. Функціональні вимоги до системи', 2)

add_paragraph(doc, '''Система повинна забезпечувати наступний функціонал:''')

reqs = [
    'Завантаження та обробка документів у форматах PDF, DOCX, XLSX.',
    'Автоматична векторизація документів з використанням embedding моделі.',
    'Зберігання векторних представлень у MongoDB Atlas Vector Search.',
    'Семантичний пошук релевантних фрагментів за запитом користувача.',
    'Генерація відповідей українською мовою на основі знайдених фрагментів.',
    'Надання посилань на джерела інформації у відповідях.',
    'REST API для інтеграції з іншими системами університету.',
    'Можливість видалення документів з бази знань.',
    'Моніторинг стану системи через health check endpoints.'
]

for req in reqs:
    p = doc.add_paragraph(req, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

add_heading(doc, '2.2. Нефункціональні вимоги', 2)

nf_reqs = {
    'Продуктивність': 'Час відповіді на запит не більше 10 секунд.',
    'Масштабованість': 'Підтримка бази знань до 10000 документів.',
    'Точність': 'Релевантність відповідей не менше 85%.',
    'Доступність': 'Час роботи системи 99% (uptime).',
    'Безпека': 'Авторизація доступу до API.',
    'Підтримка мов': 'Повна підтримка української мови.'
}

for category, req in nf_reqs.items():
    add_paragraph(doc, req, f'{category}:')

add_heading(doc, '2.3. Архітектура системи', 2)

add_paragraph(doc, '''Система побудована за мікросервісною архітектурою з розділенням на компоненти:''')

add_paragraph(doc, 'FastAPI Backend', 'Компонент 1: FastAPI Backend')
add_paragraph(doc, '''Основний веб-сервер, який обробляє HTTP запити, координує роботу інших компонентів. Реалізує REST API endpoints для завантаження документів та обробки запитань.''')

add_paragraph(doc, 'Document Parser Service', 'Компонент 2: Document Parser Service')
add_paragraph(doc, '''Відповідальний за парсинг документів різних форматів. Використовує бібліотеки PyPDF2 (PDF), python-docx (DOCX), pandas (XLSX).''')

add_paragraph(doc, 'Vector Store Service (LangChain + MongoDB Atlas)', 'Компонент 3: Vector Store Service')
add_paragraph(doc, '''Управляє векторизацією та зберіганням документів. Використовує FastEmbed для генерації embeddings та MongoDB Atlas для зберігання.''')

add_paragraph(doc, 'RAG Chain (LangChain)', 'Компонент 4: RAG Chain')
add_paragraph(doc, '''Реалізує RAG pipeline: пошук релевантних фрагментів, формування промпту, виклик LLM, повернення відповіді з джерелами.''')

add_paragraph(doc, 'LLM Integration (Deepseek)', 'Компонент 5: LLM Integration')
add_paragraph(doc, '''Інтеграція з Deepseek API для генерації відповідей. Deepseek обрано за критеріями: підтримка української мови, низька вартість, висока якість.''')

add_heading(doc, '2.4. Вибір технологій розробки', 2)

technologies = {
    'Python 3.11+': 'Основна мова програмування. Обрана за наявність потужних бібліотек для NLP та ML.',
    'FastAPI': 'Web framework для створення REST API. Переваги: висока продуктивність, автоматична документація, асинхронність.',
    'LangChain': 'Framework для RAG систем. Надає готові інтеграції з LLM та векторними БД.',
    'MongoDB Atlas': 'Хмарна БД з підтримкою Vector Search. Переваги: масштабованість, managed service, безкоштовний tier.',
    'FastEmbed': 'Бібліотека для генерації embeddings. Модель: intfloat/multilingual-e5-large (1024 dim).',
    'Deepseek': 'LLM для генерації відповідей. API-сумісний з OpenAI, низька вартість.',
    'Uvicorn': 'ASGI сервер для запуску FastAPI.',
    'Pydantic': 'Валідація даних та створення API моделей.'
}

for tech, desc in technologies.items():
    add_paragraph(doc, desc, f'{tech}:')

add_heading(doc, '2.5. Проектування структури даних', 2)

add_paragraph(doc, 'Колекція documents (метадані документів)', 'Колекція 1: documents')
add_paragraph(doc, '''Зберігає інформацію про завантажені документи: filename, file_type, uploaded_at, total_chunks, chunk_ids.''')

add_paragraph(doc, 'Колекція document_chunks (векторні представлення)', 'Колекція 2: document_chunks')
add_paragraph(doc, '''Зберігає фрагменти документів з embeddings: text (текст фрагменту), embedding (вектор 1024 dim), source_file, chunk_index, metadata.''')

add_paragraph(doc, '''Для колекції document_chunks створюється Vector Search Index з назвою vector_index, який використовує cosine similarity для пошуку схожості.''')

doc.add_page_break()

print('✓ Додано РОЗДІЛ 2')
doc.save('RAG_Report_Draft.docx')
print('✓ Документ збережено')
