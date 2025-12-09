from typing import Optional
import io
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pandas as pd


class DocumentParser:
    """Parser for different document formats"""

    @staticmethod
    async def parse_pdf(file_content: bytes) -> str:
        """
        Parse PDF file and extract text

        Args:
            file_content: PDF file content as bytes

        Returns:
            Extracted text from PDF
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text:
                    text_parts.append(f"--- Сторінка {page_num} ---\n{text}\n")

            return "\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")

    @staticmethod
    async def parse_docx(file_content: bytes) -> str:
        """
        Parse DOCX file and extract text

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Extracted text from DOCX
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        table_text.append(row_text)

                if table_text:
                    text_parts.append("\n--- Таблиця ---")
                    text_parts.extend(table_text)
                    text_parts.append("--- Кінець таблиці ---\n")

            return "\n\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")

    @staticmethod
    async def parse_xlsx(file_content: bytes) -> str:
        """
        Parse XLSX file and extract text

        Args:
            file_content: XLSX file content as bytes

        Returns:
            Extracted text from XLSX (formatted as tables)
        """
        try:
            xlsx_file = io.BytesIO(file_content)
            excel_file = pd.ExcelFile(xlsx_file)

            text_parts = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)

                text_parts.append(f"\n--- Аркуш: {sheet_name} ---")

                # Convert dataframe to readable text
                # Replace NaN with empty string
                df = df.fillna('')

                # Create header
                header = " | ".join(str(col) for col in df.columns)
                text_parts.append(header)
                text_parts.append("-" * len(header))

                # Add rows
                for _, row in df.iterrows():
                    row_text = " | ".join(str(val) for val in row.values)
                    text_parts.append(row_text)

                text_parts.append(f"--- Кінець аркуша {sheet_name} ---\n")

            return "\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Error parsing XLSX: {str(e)}")

    @classmethod
    async def parse_file(cls, file_content: bytes, file_type: str) -> str:
        """
        Parse file based on its type

        Args:
            file_content: File content as bytes
            file_type: File extension (pdf, docx, xlsx)

        Returns:
            Extracted text from file
        """
        file_type = file_type.lower()

        if file_type == "pdf":
            return await cls.parse_pdf(file_content)
        elif file_type == "docx":
            return await cls.parse_docx(file_content)
        elif file_type == "xlsx":
            return await cls.parse_xlsx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
