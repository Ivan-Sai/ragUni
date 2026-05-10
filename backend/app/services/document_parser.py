"""Document parsing for upload — extracts plain text from PDF/DOCX/XLSX/TXT.

For PDF we use pdfplumber rather than PyPDF2 because the corpus has a
lot of tabular content (exam schedules, grade sheets, course timetables)
that PyPDF2 streams as a flat text run with column boundaries lost.
pdfplumber extracts tables as 2D cell arrays and we serialise them as
Markdown so the embedding model and LLM see the row/column structure.
"""

from __future__ import annotations

import io
import logging
import re
import struct
import zipfile
from zipfile import BadZipFile

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument

# pdfplumber wraps pdfminer + a stack of helpers that can surface a
# wide variety of internal errors on malformed PDFs. We treat them all
# as "this file is broken" rather than 500ing — but we list the
# concrete exception classes here instead of catching `Exception` to
# stay within the project's "no bare except" rule.
try:
    from pdfminer.psparser import PSException  # type: ignore[import-not-found]
    from pdfminer.pdfparser import PDFSyntaxError  # type: ignore[import-not-found]
    _PDF_PARSE_EXCEPTIONS: tuple[type[BaseException], ...] = (
        PSException,
        PDFSyntaxError,
        struct.error,
        AttributeError,
        KeyError,
        IndexError,
        ValueError,
        TypeError,
        ZeroDivisionError,
        AssertionError,
    )
except ImportError:
    _PDF_PARSE_EXCEPTIONS = (
        struct.error,
        AttributeError,
        KeyError,
        IndexError,
        ValueError,
        TypeError,
        ZeroDivisionError,
        AssertionError,
    )

logger = logging.getLogger(__name__)

# Maximum page count we will attempt — guards against hostile / huge PDFs
# without an explicit upload-side check breaking parsing midway.
_MAX_PDF_PAGES = 500


# DOCX/XLSX are zip files. A 10 MB compressed payload can decompress to
# multiple GB ("zip bomb") and exhaust memory before any of our other
# guards fire. Caps below are checked against the SUM of decompressed
# member sizes BEFORE we hand the file to python-docx / openpyxl.
_MAX_OOXML_TOTAL_DECOMPRESSED: int = 100 * 1024 * 1024  # 100 MB
_MAX_OOXML_SINGLE_MEMBER: int = 50 * 1024 * 1024  # 50 MB


def _check_zip_bomb(file_content: bytes, *, kind: str) -> None:
    """Refuse OOXML payloads whose members decompress to suspicious sizes.

    Raises ``ValueError`` (mapped to HTTP 400 by the upload endpoint)
    when the sum of declared decompressed sizes exceeds the total cap
    or any single member exceeds the per-member cap. The check uses
    the ZIP central directory — no decompression happens here, so the
    guard itself cannot be DoSed by malicious input.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
            total = 0
            for info in zf.infolist():
                if info.file_size > _MAX_OOXML_SINGLE_MEMBER:
                    raise ValueError(
                        f"{kind} member {info.filename!r} would decompress "
                        f"to {info.file_size} bytes (limit "
                        f"{_MAX_OOXML_SINGLE_MEMBER})"
                    )
                total += info.file_size
                if total > _MAX_OOXML_TOTAL_DECOMPRESSED:
                    raise ValueError(
                        f"{kind} would decompress to more than "
                        f"{_MAX_OOXML_TOTAL_DECOMPRESSED} bytes — refusing"
                    )
    except BadZipFile as exc:
        raise ValueError(f"Invalid {kind} file (bad ZIP structure)") from exc


class DocumentParser:
    """Parser for different document formats."""

    @staticmethod
    async def parse_pdf(file_content: bytes) -> str:
        """Parse a PDF into a text representation that preserves tables.

        Algorithm per page:
          1. Extract every table with pdfplumber and convert each to a
             Markdown table (with merged-parent headers flattened — see
             ``_table_to_markdown``).
          2. Extract the page's plain prose, then strip out any spans
             that belong to a table we already serialised — otherwise
             every table cell would appear twice (once flat, once
             structured).
          3. Emit the page header, the prose, and the Markdown tables
             with explicit ``--- Table ---`` delimiters so the chunker
             can keep tables intact when possible.

        Tables that span multiple pages — common for Ukrainian exam
        schedules — only carry their column headers on the first page.
        We thread the last seen header set forward so a continuation
        page's table reuses them instead of falling back to col1..colN.
        """
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                if len(pdf.pages) > _MAX_PDF_PAGES:
                    raise ValueError(
                        f"PDF has {len(pdf.pages)} pages, "
                        f"maximum supported is {_MAX_PDF_PAGES}",
                    )

                page_texts: list[str] = []
                last_headers: list[str] | None = None
                for page_num, page in enumerate(pdf.pages, 1):
                    page_block, last_headers = _render_page(
                        page, page_num, last_headers
                    )
                    if page_block:
                        page_texts.append(page_block)

                return "\n".join(page_texts)

        except ValueError:
            raise
        except (OSError, UnicodeDecodeError) as exc:
            logger.warning("Error reading PDF: %s", type(exc).__name__)
            raise ValueError("Could not read PDF file.") from exc
        except _PDF_PARSE_EXCEPTIONS as exc:
            # pdfplumber can surface a wide menagerie of internal errors
            # (pdfminer.PSEOF, struct.error, ...) on malformed input —
            # treat them all as "this file is broken" rather than 500ing.
            # The exception list is enumerated at module top so that
            # MemoryError / KeyboardInterrupt / SystemExit still
            # propagate normally.
            logger.warning("pdfplumber failed: %s: %s", type(exc).__name__, exc)
            raise ValueError(
                "Could not parse PDF file. The file may be corrupted or encrypted."
            ) from exc

    @staticmethod
    async def extract_schedule_cells(file_content: bytes) -> "list[CellEvent]":
        """Extract pre-attributed schedule cells from a PDF.

        Used by the schedule extractor as the deterministic primary
        path: walks every pdfplumber table, applies vertical-text
        destacking, then asks ``schedule_table_parser`` to identify
        column→group mappings and emit one ``CellEvent`` per
        (group, day, time) slot.

        Multi-page consistency
        ----------------------

        Ukrainian week-grid PDFs print the column header (groups,
        years, levels) once on page 1 and leave pages 2/3 as raw
        data rows. Worse, pdfplumber's auto column detection often
        finds a *different* number of columns on continuation pages
        because narrower text or missing border cues let it split a
        single visual column into two. The naive "thread page 1
        columns by index" approach then mis-attributes every cell.

        Fix: extract page 1 with default settings to learn the
        canonical vertical grid (cell bbox x-coordinates), then
        re-extract pages 2..N with ``explicit_vertical_lines`` set
        to those same coordinates. All pages now report the same
        column count and ``column[i]`` means the same group on
        every page.

        Returns an empty list when no table on any page parses as a
        schedule grid; the caller should fall back to LLM-only
        extraction in that case.
        """
        from app.services.schedule_table_parser import (  # local import: avoids circular deps
            CellEvent,
            parse_schedule_table,
            parse_schedule_table_with_columns,
        )

        cells: list[CellEvent] = []
        last_columns = None
        canonical_ranges: list[tuple[float, float]] | None = None
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    # Continuation pages get re-emitted onto page 1's
                    # canonical column grid via x-overlap before
                    # parsing — preserves text content while keeping
                    # column index meaningful across pages.
                    page_rows: list[list[list[str]]] = []
                    if canonical_ranges:
                        remapped = _remap_table_to_canonical_grid(
                            page, canonical_ranges
                        )
                        if remapped:
                            page_rows.append(remapped)
                        else:
                            for table in page.extract_tables() or []:
                                page_rows.append(table)
                    else:
                        for table in page.extract_tables() or []:
                            page_rows.append(table)

                    for table in page_rows:
                        if len(table) < 2 or max(len(r) for r in table) < 3:
                            continue
                        rows = [
                            [(c or "").replace("\n", " ").strip() for c in row]
                            for row in table
                        ]
                        rows = _destack_vertical_text(rows)
                        result = parse_schedule_table(rows)
                        if result.parsed_successfully:
                            cells.extend(result.cells)
                            last_columns = result.columns
                            if canonical_ranges is None:
                                canonical_ranges = _table_canonical_x_ranges(page)
                        elif last_columns:
                            result2 = parse_schedule_table_with_columns(
                                rows, last_columns,
                            )
                            if result2.parsed_successfully:
                                cells.extend(result2.cells)
        except (OSError, UnicodeDecodeError) as exc:
            logger.warning(
                "extract_schedule_cells: pdfplumber error %s",
                type(exc).__name__,
            )
            return []
        except _PDF_PARSE_EXCEPTIONS as exc:
            logger.warning(
                "extract_schedule_cells failed (%s): %s",
                type(exc).__name__,
                exc,
            )
            return []

        logger.info(
            "extract_schedule_cells: %d pre-attributed cells",
            len(cells),
        )
        return cells

    @staticmethod
    async def parse_docx(file_content: bytes) -> str:
        """Parse DOCX file and extract text + tables.

        Refuses the file before opening it if the ZIP central directory
        announces decompressed sizes that exceed our caps — defends
        against zip-bomb DoS where a 10 MB upload would balloon to GB
        of memory.
        """
        _check_zip_bomb(file_content, kind="DOCX")
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            text_parts: list[str] = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                table_text: list[str] = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("\n--- Table ---")
                    text_parts.extend(table_text)
                    text_parts.append("--- End of table ---\n")

            return "\n\n".join(text_parts)

        except BadZipFile:
            logger.warning("Invalid DOCX file (bad ZIP structure)")
            raise ValueError("Could not parse DOCX file. The file may be corrupted.")
        except (KeyError, OSError) as exc:
            logger.warning("Error reading DOCX: %s", type(exc).__name__)
            raise ValueError("Could not read DOCX file.") from exc

    @staticmethod
    async def parse_xlsx(file_content: bytes) -> str:
        """Parse XLSX file and extract text.

        Same zip-bomb guard as ``parse_docx`` — XLSX is also OOXML and
        equally vulnerable.
        """
        _check_zip_bomb(file_content, kind="XLSX")
        try:
            xlsx_file = io.BytesIO(file_content)
            # Pin the engine so we don't silently fall through to xlrd
            # (which is read-only legacy and has its own quirks).
            excel_file = pd.ExcelFile(xlsx_file, engine="openpyxl")

            text_parts: list[str] = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text_parts.append(f"\n--- Sheet: {sheet_name} ---")

                df = df.fillna("")

                header = " | ".join(str(col) for col in df.columns)
                text_parts.append(header)
                text_parts.append("-" * len(header))
                for _, row in df.iterrows():
                    text_parts.append(" | ".join(str(val) for val in row.values))

                text_parts.append(f"--- End of sheet {sheet_name} ---\n")

            return "\n".join(text_parts)

        except BadZipFile:
            logger.warning("Invalid XLSX file (bad ZIP structure)")
            raise ValueError("Could not parse XLSX file. The file may be corrupted.")
        except (pd.errors.ParserError, pd.errors.EmptyDataError) as exc:
            logger.warning("Error parsing XLSX data: %s", type(exc).__name__)
            raise ValueError("Could not parse spreadsheet data.") from exc
        except (KeyError, OSError) as exc:
            logger.warning("Error reading XLSX: %s", type(exc).__name__)
            raise ValueError("Could not read XLSX file.") from exc

    @classmethod
    async def parse_file(cls, file_content: bytes, file_type: str) -> str:
        """Parse file based on its type."""
        file_type = file_type.lower()
        if file_type == "pdf":
            return await cls.parse_pdf(file_content)
        if file_type == "docx":
            return await cls.parse_docx(file_content)
        if file_type == "xlsx":
            return await cls.parse_xlsx(file_content)
        if file_type == "txt":
            text = file_content.decode("utf-8", errors="replace")
            if "�" in text:
                logger.warning("File contains invalid UTF-8 characters that were replaced")
            return text
        raise ValueError(f"Unsupported file type: {file_type}")


# ---------------------------------------------------------------------------
# Internal helpers — PDF page rendering
# ---------------------------------------------------------------------------


def _render_page(
    page,
    page_num: int,
    inherited_headers: list[str] | None,
) -> tuple[str, list[str] | None]:
    """Render one pdfplumber page; returns (text, headers-for-next-page)."""
    parts: list[str] = [f"--- Page {page_num} ---"]

    raw_tables = page.find_tables()
    table_bboxes = [t.bbox for t in raw_tables]

    # 1. Prose first (everything outside the table bounding boxes).
    prose = _extract_prose_outside_tables(page, table_bboxes)
    if prose.strip():
        parts.append(prose.strip())

    # 2. Tables — each as Markdown so headers + columns survive
    #    chunking and the LLM can read it as structured data.
    last_headers = inherited_headers
    for index, table_obj in enumerate(raw_tables, 1):
        rows = table_obj.extract()
        markdown, headers = _table_to_markdown(
            rows, fallback_headers=last_headers
        )
        if headers:
            last_headers = headers
        if markdown:
            parts.append(f"\n--- Table {index} ---\n{markdown}\n--- End of table ---")

    return "\n".join(parts), last_headers


def _extract_prose_outside_tables(page, table_bboxes: list[tuple]) -> str:
    """Page text minus any character that lies inside a detected table."""
    if not table_bboxes:
        return page.extract_text() or ""

    def _outside_tables(obj) -> bool:
        # `obj` is a pdfplumber character with x0/x1/top/bottom keys.
        cx = (obj["x0"] + obj["x1"]) / 2.0
        cy = (obj["top"] + obj["bottom"]) / 2.0
        for x0, top, x1, bottom in table_bboxes:
            if x0 <= cx <= x1 and top <= cy <= bottom:
                return False
        return True

    filtered = page.filter(_outside_tables)
    return filtered.extract_text() or ""


# Recognisable Ukrainian "marker" words that appear vertically in
# university documents. Used to validate the orientation of a
# de-stacked column — if the assembled string matches one of these
# (or a substring thereof), we assume the run was a real label.
_VERTICAL_MARKER_WORDS: tuple[str, ...] = (
    "понеділок", "вівторок", "середа", "четвер", "пятниця", "п'ятниця",
    "субота", "неділя",
    "години", "год", "час",
    "дні", "день", "тиждень",
    "група", "групи",
)


def _normalise_for_match(s: str) -> str:
    """Strip apostrophes / spaces / diacritics so vertically-stacked
    words compare equal to their dictionary form."""
    out = s.lower().replace(" ", "").replace("'", "").replace("’", "")
    return out


def _match_one_marker(chars: list[str]) -> tuple[str, int] | None:
    """Look for ONE marker word in ``chars``, trying forward and
    reverse orientation. Return ``(label, length_consumed)`` for the
    LONGEST matching marker, or ``None`` if no marker fits.

    Used by ``_destack_vertical_text`` to peel a single day name off
    a long run that may contain several stacked words back-to-back
    (e.g. ``вівторок`` immediately followed by ``середа`` in the
    same column on a multi-day continuation page).
    """
    if not chars:
        return None

    def _raw_consumed(orientation: list[str], marker_norm: str) -> int | None:
        """Return how many raw chars in ``orientation`` correspond to
        the leading ``marker_norm`` in normalised form, or ``None`` if
        the orientation does not start with the marker.

        Walks ``orientation`` one raw char at a time and counts only
        chars that survive normalisation (apostrophes are dropped by
        ``_normalise_for_match``). When the running normalised buffer
        equals ``marker_norm``, the current raw offset is returned —
        so ``"п’ятниця"`` with marker ``"пятниця"`` returns 8, not 7.
        """
        norm_seen = ""
        for raw_idx, raw_char in enumerate(orientation, start=1):
            norm_seen += _normalise_for_match(raw_char)
            if not marker_norm.startswith(norm_seen):
                return None
            if norm_seen == marker_norm:
                return raw_idx
        return None

    # Try forward (top-down) and reverse (bottom-up) orientations.
    # Prefer the longer match so "понеділок" wins over "ділок".
    candidates: list[tuple[str, int]] = []
    sorted_markers = sorted(_VERTICAL_MARKER_WORDS, key=len, reverse=True)
    for orientation in (chars, list(reversed(chars))):
        for marker in sorted_markers:
            marker_norm = _normalise_for_match(marker)
            if not marker_norm:
                continue
            consumed = _raw_consumed(orientation, marker_norm)
            if consumed is None:
                continue
            label = "".join(orientation[:consumed])
            candidates.append((label, consumed))
            break
    if not candidates:
        return None
    return max(candidates, key=lambda c: c[1])


def _pick_destacked_orientation(chars_top_down: list[str]) -> str | None:
    """Backwards-compatible wrapper around ``_match_one_marker`` —
    returns just the label when the run matches a marker exactly
    (whole-run match). Used by the within-cell destack helper.
    """
    if len(chars_top_down) < 3:
        return None
    match = _match_one_marker(chars_top_down)
    if not match:
        return None
    label, consumed = match
    if consumed != len(chars_top_down):
        return None
    return label


def _table_canonical_x_ranges(page) -> list[tuple[float, float]] | None:
    """Return the ``(x0, x1)`` range of every column in the first
    table on ``page``.

    Used by ``extract_schedule_cells`` to lock page 1's column
    boundaries and remap subsequent pages' cells onto the same
    column index by x-overlap. Returns ``None`` when the page has
    no recognisable table.
    """
    try:
        tables = page.find_tables()
    except _PDF_PARSE_EXCEPTIONS:
        return None
    if not tables:
        return None
    table = tables[0]
    xs: set[float] = set()
    for cell in table.cells or []:
        if cell:
            x0, _top, x1, _bottom = cell
            xs.add(round(float(x0), 1))
            xs.add(round(float(x1), 1))
    if len(xs) < 3:
        return None
    sorted_xs = sorted(xs)
    return list(zip(sorted_xs[:-1], sorted_xs[1:]))


def _remap_table_to_canonical_grid(
    page,
    canonical_ranges: list[tuple[float, float]],
    canonical_y_lines: list[float] | None = None,
) -> list[list[str]] | None:
    """Re-emit a page's content onto page 1's canonical column grid.

    Walks ``Table.cells`` (the original physical cell bboxes from
    pdfplumber) and, for every cell, places its text in *every*
    canonical column the cell horizontally overlaps. Joint cells —
    where one entry visually spans two adjacent group columns — are
    therefore preserved on both groups instead of being shredded
    into character fragments.

    Each cell's text is read by cropping the page to the cell bbox
    and running ``extract_text``; row indexing is done by clustering
    cells by their ``top`` y-coordinate. The result has exactly the
    same column count as page 1's canonical grid, so downstream
    parsing (day/time/group attribution) sees a consistent shape on
    every page.

    Falls back to a word-level x-centre map when the page has no
    detectable table (for sparse pages where ``find_tables`` returns
    empty), so a misdetected page never silently disappears.
    """
    n_canonical = len(canonical_ranges)
    if n_canonical == 0:
        return None

    try:
        tables = page.find_tables()
    except _PDF_PARSE_EXCEPTIONS:
        tables = []
    if not tables:
        return _remap_words_to_canonical_grid(page, canonical_ranges)

    table = tables[0]
    rows_obj = getattr(table, "rows", None)
    if not rows_obj:
        return _remap_words_to_canonical_grid(page, canonical_ranges)

    # Extract words once for the whole page. Word centres tell us
    # which physical cell a word belongs to without ever cropping —
    # cropping picks up glyphs from neighbouring cells when the
    # crop padding is even slightly generous, which produced
    # characters-interleaved garbage on dense rows.
    try:
        page_words = page.extract_words(
            use_text_flow=False,
            keep_blank_chars=False,
        )
    except _PDF_PARSE_EXCEPTIONS:
        page_words = []

    def _word_center(w: dict) -> tuple[float, float]:
        return (
            (float(w["x0"]) + float(w["x1"])) / 2.0,
            (float(w["top"]) + float(w["bottom"])) / 2.0,
        )

    def _overlapping_cols(x0: float, x1: float) -> list[int]:
        """All canonical column indices whose x-range overlaps the
        cell's x-range by at least 1pt."""
        out: list[int] = []
        for idx, (rx0, rx1) in enumerate(canonical_ranges):
            overlap = min(x1, rx1) - max(x0, rx0)
            if overlap >= 1.0:
                out.append(idx)
        return out

    def _cell_text(bbox: tuple[float, float, float, float]) -> str:
        """Assemble the cell's text from words whose centre lies
        strictly inside the cell bbox. Words are emitted in
        natural reading order (top-to-bottom, then left-to-right
        within each line)."""
        x0, top, x1, bottom = bbox
        contained = []
        for w in page_words:
            cx, cy = _word_center(w)
            if x0 <= cx <= x1 and top <= cy <= bottom:
                contained.append(w)
        if not contained:
            return ""
        contained.sort(key=lambda w: (round(float(w["top"]), 1), float(w["x0"])))
        # Group by visual line (≈3.5pt y-tolerance) and join words
        # within a line with single spaces, lines with single spaces
        # too — the schedule parser doesn't care about line breaks.
        out_parts: list[str] = []
        current_line: list[str] = []
        current_top: float | None = None
        for w in contained:
            top_w = float(w["top"])
            if current_top is None or abs(top_w - current_top) <= 3.5:
                current_line.append(w.get("text", ""))
                current_top = top_w if current_top is None else current_top
            else:
                if current_line:
                    out_parts.append(" ".join(current_line))
                current_line = [w.get("text", "")]
                current_top = top_w
        if current_line:
            out_parts.append(" ".join(current_line))
        return " ".join(p for p in out_parts if p).strip()

    remapped: list[list[str]] = []
    for row_obj in rows_obj:
        row_cells_raw = [
            (float(c[0]), float(c[1]), float(c[2]), float(c[3]))
            for c in (row_obj.cells or [])
            if c is not None
        ]
        if not row_cells_raw:
            continue
        # Sort the row's cells left-to-right.
        row_cells_raw.sort(key=lambda c: c[0])

        out_row = [""] * n_canonical
        for x0, top, x1, bottom in row_cells_raw:
            text = _cell_text((x0, top, x1, bottom))
            if not text:
                continue
            for col_idx in _overlapping_cols(x0, x1):
                if out_row[col_idx]:
                    out_row[col_idx] += " " + text
                else:
                    out_row[col_idx] = text
        if any(out_row):
            remapped.append(out_row)
    return remapped


def _remap_words_to_canonical_grid(
    page,
    canonical_ranges: list[tuple[float, float]],
) -> list[list[str]] | None:
    """Word-level fallback used when ``find_tables`` returns nothing.

    Each word is dropped into the canonical column whose x-range
    covers the word's centre. Cannot recover horizontally-merged
    joint cells — that's why the cell-bbox path above is preferred —
    but is good enough for sparse pages that have no detectable
    table grid.
    """
    try:
        words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
    except _PDF_PARSE_EXCEPTIONS:
        return None
    if not words:
        return None

    n_canonical = len(canonical_ranges)

    sorted_words = sorted(words, key=lambda w: (round(w["top"], 1), w["x0"]))
    rows: list[list[dict]] = []
    current_top: float | None = None
    current_row: list[dict] = []
    for w in sorted_words:
        top = float(w["top"])
        if current_top is None or abs(top - current_top) <= 3.5:
            current_row.append(w)
            current_top = top if current_top is None else current_top
        else:
            rows.append(current_row)
            current_row = [w]
            current_top = top
    if current_row:
        rows.append(current_row)

    def _assign(x0: float, x1: float) -> int | None:
        centre = (x0 + x1) / 2.0
        for idx, (rx0, rx1) in enumerate(canonical_ranges):
            if rx0 - 0.5 <= centre <= rx1 + 0.5:
                return idx
        return None

    remapped: list[list[str]] = []
    for word_row in rows:
        out_row = [""] * n_canonical
        for w in word_row:
            target = _assign(float(w["x0"]), float(w["x1"]))
            if target is None:
                continue
            text = w.get("text", "")
            if not text:
                continue
            if out_row[target]:
                out_row[target] += " " + text
            else:
                out_row[target] = text
        if any(out_row):
            remapped.append(out_row)
    return remapped


# Within-cell vertical stack: a single cell whose content is
# letter-by-letter separated by whitespace (e.g. "к о л і д е н о п"
# = "понеділок" reversed, or "я ц и н т я ’ п" = "п'ятниця"). The
# pattern allows apostrophes — Ukrainian names like п'ятниця print
# them as separate stack elements.
_STACK_CHAR_CLASS = r"(?:[^\W\d_]|[''’`])"
_WITHIN_CELL_STACK_RE = re.compile(
    rf"^{_STACK_CHAR_CLASS}(?:\s+{_STACK_CHAR_CLASS}){{3,}}$",
    flags=re.UNICODE,
)


def _within_cell_destack(cell: str) -> str:
    """Resolve a within-cell vertical stack to its proper word.

    Returns the cell unchanged when the content does not match the
    space-separated single-letter pattern, or when the assembled
    word does not look like a known marker. Conservative on
    purpose — this runs on every cell, so a false positive would
    eat real data.
    """
    stripped = cell.strip()
    if not _WITHIN_CELL_STACK_RE.match(stripped):
        return cell
    chars = stripped.split()
    label = _pick_destacked_orientation(chars)
    return label if label else cell


def _destack_vertical_text(rows: list[list[str]]) -> list[list[str]]:
    """Collapse vertically-written labels into a single cell per run.

    Two patterns are handled:

    * Across-row stack — many consecutive rows of single-char cells
      in the same column. Common when pdfplumber treats each line
      of vertical text as its own row.
    * Within-cell stack — a single cell whose content is the
      vertical word written as space-separated letters. Common when
      pdfplumber merges the visual stack into one logical cell.

    Both end up as the assembled word in the first relevant cell
    so the LLM extractor never sees the raw stack.
    """
    if not rows:
        return rows

    width = max(len(row) for row in rows)
    result = [list(row) + [""] * (width - len(row)) for row in rows]

    # Pass 1: within-cell destacking. Each cell handled independently
    # so this is order-free and idempotent.
    for r_idx in range(len(result)):
        for c_idx in range(width):
            result[r_idx][c_idx] = _within_cell_destack(result[r_idx][c_idx])

    for col_idx in range(width):
        run_indices: list[int] = []
        run_chars: list[str] = []

        def flush_run() -> None:
            nonlocal run_indices, run_chars
            try:
                # Pair up indices with their letter; drop the empties.
                paired = [
                    (idx, char)
                    for idx, char in zip(run_indices, run_chars)
                    if char
                ]
                if len(paired) < 4:
                    return
                # Split the run at large row-index gaps. A real
                # vertical word stack is compact — each letter sits
                # one or two rows below the previous. A bigger gap
                # (≥6 rows of empty cells) almost always means we
                # crossed into the next day's stack on a continuation
                # page that stacks several days in the same column.
                gap_threshold = 6
                groups: list[list[tuple[int, str]]] = [[paired[0]]]
                for prev, curr in zip(paired, paired[1:]):
                    if curr[0] - prev[0] >= gap_threshold:
                        groups.append([curr])
                    else:
                        groups[-1].append(curr)

                for group in groups:
                    if len(group) < 4:
                        continue
                    chars_only = [c for _, c in group]
                    label = _pick_destacked_orientation(chars_only)
                    if not label:
                        continue
                    anchor = group[0][0]
                    result[anchor][col_idx] = label
                    for idx, _ in group[1:]:
                        result[idx][col_idx] = ""
            finally:
                run_indices = []
                run_chars = []

        for row_idx in range(len(result)):
            cell = result[row_idx][col_idx].strip()
            # A "stack candidate" is at most one Unicode word char.
            if len(cell) <= 1 and (not cell or cell.isalpha()):
                run_indices.append(row_idx)
                run_chars.append(cell)
                continue
            # Multiple short tokens separated by whitespace — common
            # when pdfplumber's word splitter merges two adjacent
            # vertical-stack rows (e.g. "’ п" for the apostrophe and
            # п at the bottom of "п'ятниця"). Treat each whitespace-
            # separated single-glyph token as if it were its own row
            # entry so the run picks up the trailing letters.
            tokens = cell.split()
            if tokens and all(
                len(t) == 1 and (t.isalpha() or t in "'’`")
                for t in tokens
            ):
                for t in tokens:
                    run_indices.append(row_idx)
                    run_chars.append(t)
                continue
            flush_run()
        flush_run()

    return result


def _table_to_markdown(
    rows: list[list[str | None]],
    fallback_headers: list[str] | None = None,
) -> tuple[str, list[str] | None]:
    """Serialise a table to Markdown.

    Returns ``(markdown, headers)`` where ``headers`` is the list used
    for this table's columns — callers thread it forward to the next
    page so a continuation table inherits the parent's column names
    instead of devolving to ``col1..colN``.

    Handles the merged-parent-header pattern common in Ukrainian
    university schedules — e.g. a single "Екзамени" cell spanning
    "Дата | Год. | Ауд." beneath.

    Also collapses vertically-written labels: timetables print day
    names as one character per cell down the leftmost column ("П", "О",
    "Н", "Е", "Д", …). pdfplumber returns them as separate cells, so
    the LLM extractor sees the day spelled letter-by-letter and never
    captures it. ``_destack_vertical_text`` reassembles those columns.
    """
    cleaned: list[list[str]] = [
        [(cell or "").replace("\n", " ").strip() for cell in row]
        for row in rows
        if row and any((cell or "").strip() for cell in row)
    ]
    if not cleaned:
        return "", None

    width = max(len(r) for r in cleaned)
    cleaned = [row + [""] * (width - len(row)) for row in cleaned]
    cleaned = _destack_vertical_text(cleaned)

    header_height = _detect_header_height(cleaned)
    if header_height == 0:
        if fallback_headers and len(fallback_headers) == width:
            headers = fallback_headers
        else:
            headers = [f"col{i + 1}" for i in range(width)]
        body = cleaned
    else:
        headers = _merge_headers(cleaned[:header_height], width)
        body = cleaned[header_height:]

    if not body:
        return "", headers

    md_lines: list[str] = []
    md_lines.append("| " + " | ".join(headers) + " |")
    md_lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in body:
        md_lines.append("| " + " | ".join(cell or " " for cell in row) + " |")

    return "\n".join(md_lines), headers


def _looks_like_header_cell(cell: str) -> bool:
    """Heuristic: header cells are short pure-text labels.

    Empty cells are considered header-compatible because merged spans
    legitimately leave the lower-row sub-cells blank. A cell containing
    digits or longer-than-label text (>40 chars) is treated as data —
    that's how we tell a "Дата | Год. | Ауд." header row apart from a
    "13.05.2026 | 12:00 | ауд. 8" data row.
    """
    s = cell.strip()
    if not s:
        return True
    if len(s) > 40:
        return False
    if any(ch.isdigit() for ch in s):
        return False
    return True


def _detect_header_height(rows: list[list[str]]) -> int:
    """Return how many top rows form the table header (0..3).

    Walks down from the top until it finds the first row that does not
    look like a header. Returns 0 when row 0 is already data — common
    for tables that span pages and so don't repeat their header on
    continuation pages.
    """
    for i, row in enumerate(rows):
        if not all(_looks_like_header_cell(c) for c in row):
            return i
    return min(len(rows), 3)


def _merge_headers(header_rows: list[list[str]], width: int) -> list[str]:
    """Combine 0..3 header rows into a flat list of column names.

    The top row (if any) is forward-filled across empty cells so a
    "Консультації" parent that visually spans three sub-columns
    actually qualifies all three of them. Sub-labels from lower header
    rows are joined with spaces and wrapped in parentheses next to the
    parent, e.g. ``Дата (консультації)``.
    """
    if not header_rows:
        return [f"col{i + 1}" for i in range(width)]

    # Forward-fill the parent (top) row only — it's the row that
    # carries merged spans like Консультації / Екзамени.
    parents_raw = header_rows[0] + [""] * (width - len(header_rows[0]))
    filled_parents: list[str] = []
    last = ""
    for i in range(width):
        value = (parents_raw[i] or "").strip()
        if value:
            last = value
            filled_parents.append(value)
        else:
            filled_parents.append(last)

    # Lower header rows: concatenate their non-empty cells per column.
    sub_rows = header_rows[1:]
    out: list[str] = []
    for i in range(width):
        parts: list[str] = []
        for row in sub_rows:
            if i < len(row):
                cell = (row[i] or "").strip()
                if cell:
                    parts.append(cell)
        sub = " ".join(parts)
        parent = filled_parents[i]
        if sub and parent and sub.lower() != parent.lower():
            out.append(f"{sub} ({parent.lower()})")
        elif sub:
            out.append(sub)
        elif parent:
            out.append(parent)
        else:
            out.append(f"col{i + 1}")
    return out
